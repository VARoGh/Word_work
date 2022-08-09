from docx import Document
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
import re

file_name = 'ПЗ.docx'
# создание или открытие документа
doc = Document(file_name)

for i, paragraph in enumerate(doc.paragraphs):
    strings = paragraph.text
    if strings == '':
        continue

    p_fmt = paragraph.paragraph_format

    # настройка отступа обзаца слева
    p_fmt.left_indent = Mm(0)
    # настройка отступа обзаца справа
    p_fmt.right_indent = Mm(0)

    # отступ (первой) красной строки
    p_fmt.first_line_indent = Mm(12.5)

    # интервалы между абзацами - до и после
    p_fmt.space_before = Pt(0)  # Pt(12)
    p_fmt.space_after = Pt(0)  # Pt(12)

    # межстрочный интервал
    p_fmt.line_spacing = 1.5


    def run_font_bold(par, flag):
        """ Для жирного шрифта """
        for run in par.runs:
            run.font.bold = flag
        return True

    # установка высоты букв Pt(size)
    def run_font_size(par, size):
        for run in par.runs:
            if run.font.size != Pt(size):
                run.font.size = Pt(size)

    # форматирование шрифтов
    for run in paragraph.runs:
        if run.font.name != 'Times New Roman':
            run.font.name = 'Times New Roman'

    # выравнивание абзаца по ширине (), жирный и по высоте букв
    if paragraph.text in ['Введение', 'Содержание', 'Заключение', 'Перечеь использованных информационных ресурсов'] \
            or re.match(r'[0-9]\s\w+', paragraph.text):
        # paragraph.style.name == 'Heading 1'
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_font_bold(paragraph, False)
        run_font_size(paragraph, 16)
    else:
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_font_size(paragraph, 14)

    # Настройка заголовков и подзаголовков
    if re.match(r'[0-9]*\s\w+', paragraph.text) \
            or re.match(r'[0-9]*\.[0-9]\s\w+', paragraph.text) \
            or re.match(r'[0-9]*\.[0-9]*\.[0-9]*\s\w+', paragraph.text):
        # paragraph.style.name = 'Heading 1'
        p_fmt.line_spacing = 1.0
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_font_bold(paragraph, True)
        if re.match(r'\.*', doc.paragraphs[i+1].text):
            paragraph.add_run('\n')

    # Заголовки таблиц
    if re.match(r'Таблица \d', paragraph.text):
        p_fmt.line_spacing = 1.0
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_fmt.left_indent = Mm(27.5) # настройка отступа обзаца слева
        p_fmt.right_indent = Mm(0) # настройка отступа обзаца справа
        p_fmt.first_line_indent = Mm(-27.5) # отступ (первой) красной строки

    # Заголовки рисунков
    if re.match(r'Рисунок \d', paragraph.text):
        p_fmt.line_spacing = 1.0
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if re.match(r'\.*', doc.paragraphs[i+1].text):
            paragraph.add_run('\n')

doc.save(file_name)
